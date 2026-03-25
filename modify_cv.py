import sys
import subprocess

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import docx
except ImportError:
    print("Installing python-docx...")
    install("python-docx")
    import docx

doc_path = r"C:\Users\vikas\OneDrive\Desktop\vikas cv.docx"

try:
    doc = docx.Document(doc_path)
except Exception as e:
    print(f"Error opening document: {e}")
    sys.exit(1)

# 1. Update text content
for p in doc.paragraphs:
    if "aspiring software developer" in p.text.lower():
        # Let's replace the sentence or phrase
        # Specifically targeting "aspiring software developer and engineer"
        # We will remove it or change it to "Full Stack Developer"
        p.text = p.text.replace("aspiring software developer and engineer", "Full Stack Developer")
        p.text = p.text.replace("Aspiring software developer and engineer", "Full Stack Developer")
        p.text = p.text.replace("Aspiring Software Developer and Engineer", "Full Stack Developer")
        p.text = p.text.replace("aspiring software developer", "Full Stack Developer")
        p.text = p.text.replace("Aspiring software developer", "Full Stack Developer")
    
    # 2. Update 12 percentage
    if "68.2" in p.text:
        p.text = p.text.replace("68.2%", "85.2%")
        p.text = p.text.replace("68.2", "85.2")

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if "aspiring software developer" in p.text.lower():
                    p.text = p.text.replace("aspiring software developer and engineer", "Full Stack Developer")
                    p.text = p.text.replace("Aspiring software developer and engineer", "Full Stack Developer")
                    p.text = p.text.replace("Aspiring Software Developer and Engineer", "Full Stack Developer")
                    p.text = p.text.replace("aspiring software developer", "Full Stack Developer")
                    p.text = p.text.replace("Aspiring software developer", "Full Stack Developer")
                
                if "68.2" in p.text:
                    p.text = p.text.replace("68.2%", "85.2%")
                    p.text = p.text.replace("68.2", "85.2")

# 3. Add a new section at the end for theory
doc.add_heading('Technical Philosophy & Approach', level=1)
doc.add_paragraph(
    "As a Full Stack Developer, my primary objective is to build scalable, secure, and user-centric applications. "
    "I believe that great software is built on a foundation of clean code, robust architecture, and seamless user experiences. "
    "Whether I am developing a high-performance backend system using Java and Spring Boot, or crafting responsive interfaces with React, "
    "my focus remains on solving complex problems through elegant technical solutions."
)
doc.add_paragraph(
    "I approach every project with a strong emphasis on system design and software architecture. "
    "By utilizing best practices in database management and building RESTful APIs, I ensure that data flow is efficient and secure."
)

try:
    doc.save(doc_path)
    print("CV successfully modified.")
except Exception as e:
    print(f"Error saving document: {e}")
